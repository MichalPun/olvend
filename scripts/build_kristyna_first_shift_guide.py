from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "assets" / "kristyna-first-shift"
OLD_SOURCE = ROOT / "docs" / "assets" / "mobile-manual"
GENERATED = SOURCE / "generated"
OUT_DOCX = ROOT / "docs" / "OLVEND_prvni_smena_Kristyna_Dvorakova_2026-08-17.docx"

RED = "D71920"
DARK = "17202B"
MUTED = "667085"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FDEBEC"
LIGHT_GREEN = "EAF7EF"
LIGHT_AMBER = "FFF4D6"
GREEN = "16804A"
AMBER = "9A6700"
WHITE = "FFFFFF"


def rgb(value):
    return RGBColor.from_string(value)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        prevent_split(row)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "OLVEND | První směna Kristýny Dvořákové | 17. 8. 2026"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = rgb(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Interní pracovní návod | OLVEND 1.9 | Strana ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_heading(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = rgb(DARK)
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.paragraph_format.space_after = Pt(10)
        p2.runs[0].font.size = Pt(10.5)
        p2.runs[0].font.color.rgb = rgb(MUTED)


def add_callout(doc, label, text, kind="info"):
    fill, accent = {
        "stop": (LIGHT_RED, "A61B1B"),
        "ok": (LIGHT_GREEN, GREEN),
        "warn": (LIGHT_AMBER, AMBER),
        "info": (LIGHT_BLUE, "1F4D78"),
    }[kind]
    table = doc.add_table(rows=1, cols=2)
    table_geometry(table, [1.15, 5.35])
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), fill)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = left.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = rgb(WHITE)
    right = table.cell(0, 1).paragraphs[0]
    right.paragraph_format.space_after = Pt(0)
    r = right.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_step(doc, number, title, text, strong=None):
    table = doc.add_table(rows=1, cols=2)
    table_geometry(table, [0.52, 5.98])
    shade(table.cell(0, 0), RED)
    shade(table.cell(0, 1), "F6F8FA")
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = left.add_run(str(number))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = rgb(WHITE)
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(9.5)
    if strong:
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        r = p3.add_run(strong)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = rgb(GREEN)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_checklist(doc, items):
    table = doc.add_table(rows=1, cols=3)
    table_geometry(table, [0.45, 4.8, 1.25])
    for idx, head in enumerate(("", "KROK", "HOTOVO")):
        shade(table.cell(0, idx), LIGHT_BLUE)
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(head)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(MUTED)
    set_repeat_header(table.rows[0])
    for idx, text in enumerate(items, 1):
        cells = table.add_row().cells
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cells[0].paragraphs[0].add_run(str(idx))
        r.bold = True
        r.font.color.rgb = rgb(RED)
        cells[1].paragraphs[0].add_run(text).font.size = Pt(9.5)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].add_run("□").font.size = Pt(16)
    return table


def crop_phone(src, out, box=(370, 0, 910, 720), marks=None):
    image = Image.open(src).convert("RGB").crop(box)
    if marks:
        draw = ImageDraw.Draw(image)
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 22)
        except OSError:
            font = ImageFont.load_default()
        for number, x, y in marks:
            draw.ellipse((x - 18, y - 18, x + 18, y + 18), fill="#D71920", outline="#FFFFFF", width=3)
            label = str(number)
            box_text = draw.textbbox((0, 0), label, font=font)
            draw.text((x - (box_text[2] - box_text[0]) / 2, y - (box_text[3] - box_text[1]) / 2 - 2), label, font=font, fill="#FFFFFF")
    image.save(out, quality=92)
    return out


def prepare_images():
    GENERATED.mkdir(parents=True, exist_ok=True)
    output = {}
    specs = {
        "today": (SOURCE / "02-dnes.png", [(1, 465, 683)]),
        "stock": (SOURCE / "04-sklad.png", [(1, 475, 683)]),
        "loading": (SOURCE / "05-ranni-nakladka.png", [(1, 446, 486)]),
        "loading_detail": (SOURCE / "06-nakladka-mnozstvi.png", [(1, 445, 512), (2, 469, 681)]),
        "route": (SOURCE / "03-trasa.png", [(1, 461, 683)]),
        "stop": (SOURCE / "07-detail-zastavky.png", [(1, 464, 172)]),
    }
    for key, (src, marks) in specs.items():
        out = GENERATED / f"{key}.jpg"
        output[key] = crop_phone(src, out, marks=marks)
    start_src = OLD_SOURCE / "03-zahajeni-smeny.png"
    start = Image.open(start_src).convert("RGB")
    output["start"] = GENERATED / "start.jpg"
    start.save(output["start"], quality=92)
    return output


def add_image_with_caption(doc, path, caption, width=3.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    cp.runs[0].font.size = Pt(8.5)
    cp.runs[0].font.color.rgb = rgb(MUTED)


def add_two_column_visual(doc, image_path, steps):
    table = doc.add_table(rows=1, cols=2)
    table_geometry(table, [3.0, 3.5])
    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(2.75))
    right.text = ""
    for number, title, text in steps:
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{number}. {title}")
        r.bold = True
        r.font.size = Pt(10.5)
        p2 = right.add_paragraph(text)
        p2.paragraph_format.space_after = Pt(8)
        p2.runs[0].font.size = Pt(9.5)
        p2.runs[0].font.color.rgb = rgb(MUTED)


def page_break(doc):
    doc.add_page_break()


def build():
    images = prepare_images()
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("PRVNÍ SMĚNA · RYCHLÝ NÁVOD")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(RED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Kristýna Dvořáková")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = rgb(DARK)
    p = doc.add_paragraph("Pondělí 17. 8. 2026 · Fiat Doblo 5AP9000 · OLVEND 1.9")
    p.paragraph_format.space_after = Pt(15)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = rgb(MUTED)

    add_callout(doc, "Nejdůležitější", "Aplikace tě vede krok za krokem. Nic nepřeskakuj a nepotvrzuj krok, který jsi fyzicky neudělala.", "ok")
    add_heading(doc, "Celá směna v 10 krocích")
    add_checklist(doc, [
        "Otevřít OLVEND a přečíst všechny denní pokyny.",
        "Ve Skladu převzít zboží, které nachystal skladník.",
        "Zkontrolovat kusy a potvrdit převzetí vychystaného zboží.",
        "Udělat ranní nakládku baget a potvrdit skutečné množství i expirace.",
        "Vybrat Fiat Doblo 5AP9000 a zapsat skutečný počáteční stav km.",
        "Zahájit směnu a otevřít Trasu.",
        "Jet zastávky v pořadí aplikace; nejetí zastávky pouze přes žádost.",
        "U každého automatu dokončit všechny kroky a teprve potom odjet.",
        "Po návratu vyložit vratky/odpisy podle aplikace.",
        "Zapsat konečný stav km a ukončit směnu.",
    ])
    add_callout(doc, "Když si nejsi jistá", "Zastav práci, nic nehádej a zavolej vedoucímu. Chybný údaj nepotvrzuj jen proto, abys mohla pokračovat.", "warn")

    page_break(doc)
    add_heading(doc, "1. Nejdřív pokyny a převzetí směny", "Toto udělej ještě před odjezdem ze skladu.")
    add_two_column_visual(doc, images["start"], [
        ("1", "Přečti denní pokyny", "Pokyn otevři, přečti celý a potvrď až tehdy, když mu rozumíš."),
        ("2", "Vyber správné auto", "Pro zítřek zvol Fiat Doblo · SPZ 5AP9000."),
        ("3", "Zapiš start km", "Opiš skutečný stav přímo z tachometru. Nezaokrouhluj a neodhaduj."),
        ("4", "Zahaj směnu", "Tlačítko použij až po převzetí vychystaného zboží a ranní nakládce."),
    ])
    add_callout(doc, "Pozor", "Denní pokyn může blokovat zahájení směny. Není to chyba aplikace: pokyn je potřeba otevřít, přečíst a potvrdit.", "warn")

    page_break(doc)
    add_heading(doc, "2. Povinně převezmi zboží od skladníka", "Toto je jiné než ranní nakládka baget.")
    add_two_column_visual(doc, images["stock"], [
        ("1", "Otevři Sklad", "Dole klepni na Sklad."),
        ("2", "Najdi Čeká na tebe", "Když skladník něco nachystal, zobrazí se karta Převzít vychystané / Potvrdit převzetí zboží."),
        ("3", "Fyzicky zkontroluj", "Porovnej produkt a počet s tím, co opravdu nakládáš do auta."),
        ("4", "Potvrď převzetí zboží", "Bez tohoto tlačítka zůstane zboží v systému jako nepřevzaté."),
    ])
    add_callout(doc, "Stop", "Nesouhlasí produkt nebo počet? Nepotvrzuj převzetí. Zboží dej stranou a zavolej skladníkovi nebo vedoucímu.", "stop")

    page_break(doc)
    add_heading(doc, "3. Ranní nakládka chlazeného zboží", "Po převzetí zboží od skladníka otevři samostatnou ranní nakládku.")
    add_two_column_visual(doc, images["loading_detail"], [
        ("1", "Sklad → Naložit do auta", "Otevři Naložit do auta. Aplikace ukáže doporučení pro dnešní trasu."),
        ("2", "Použij doporučení", "Tlačítkem načti návrh. Potom uprav pouze podle toho, co skutečně bereš."),
        ("3", "Zkontroluj každý obal", "U každé bagety porovnej datum v aplikaci s datem na fyzickém obalu."),
        ("4", "Potvrď skutečnou nakládku", "Potvrď jen fyzicky naložené kusy. Teprve po potvrzení je nakládka hotová."),
    ])
    add_callout(doc, "Důležité", "Ranní nakládka nenahrazuje převzetí vychystaného zboží od skladníka. Zítra musí být hotové oba kroky.", "info")

    page_break(doc)
    add_heading(doc, "4. Zítřejší bagety s krátkou expirací", "Tyto kusy mají přednost. Vezmi je a dej je do určených automatů.")
    add_callout(doc, "Povinně", "Celkem 8 kusů s krátkou expirací. Nezaměň je za novější kusy a nenechávej je ve skladu ani v autě.", "stop")
    table = doc.add_table(rows=1, cols=3)
    table_geometry(table, [2.55, 1.75, 2.2])
    for idx, text in enumerate(("KAM", "KOLIK", "CO")):
        shade(table.cell(0, idx), LIGHT_BLUE)
        r = table.cell(0, idx).paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = rgb(MUTED)
    rows = [
        ("Vitar Tišnov · EV 78", "6 ks", "2× Labužník\n2× Debrecínská\n2× Trhané vepřové"),
        ("NTS Brno-Slatina · EV 65", "2 ks", "1× Kuře teriyaki\n1× Debrecínská"),
    ]
    for where, count, items in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(where).bold = True
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cells[1].paragraphs[0].add_run(count)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = rgb(RED)
        cells[2].paragraphs[0].add_run(items)
    doc.add_paragraph()
    add_step(doc, 1, "Ve skladu", "Vezmi nejdřív kusy s nejkratší expirací a fyzicky zkontroluj datum na každém obalu.")
    add_step(doc, 2, "V ranní nakládce", "Potvrď expirace a skutečně naložené množství.")
    add_step(doc, 3, "U automatu", "Vlož krátkou expiraci před novější zboží tak, aby se prodala jako první.")
    add_step(doc, 4, "Před odjezdem", "Zkontroluj, že je všech 8 kusů opravdu vložených v určených automatech.")

    page_break(doc)
    add_heading(doc, "5. Trasa: vždy dělej právě zobrazenou zastávku", "Pořadí určuje aplikace podle plánu a aktuální situace.")
    add_two_column_visual(doc, images["route"], [
        ("1", "Otevři Trasu", "Dole klepni na Trasa. Nahoře uvidíš další zastávku a plánovaný čas."),
        ("2", "Jeď podle pořadí", "Zastávky ručně nepřehazuj. Po dokončení se další pořadí samo přepočítá."),
        ("3", "Nejetí jen přes žádost", "Když místo opravdu nejde obsloužit, odešli žádost s pravdivým důvodem."),
        ("4", "Otevři automat až na místě", "Po příjezdu otevři detail zastávky a postupuj od prvního nedokončeného kroku."),
    ])
    add_callout(doc, "Nikdy", "Neoznačuj zastávku jako hotovou, pokud jsi na místě nebyla nebo nejsou dokončené všechny povinné kroky.", "stop")

    page_break(doc)
    add_heading(doc, "6. U automatu: dokonči všechny čtyři kroky", "Aplikace tě vede v pořadí Pick list → Doplnění → Hotovost → Kontrola.")
    add_two_column_visual(doc, images["stop"], [
        ("1", "Pick list", "Vezmi z auta přesně zobrazené produkty. Zkontroluj název, kusy a expiraci."),
        ("2", "Doplnění", "Zapiš skutečný stav a skutečně doplněné kusy. Prošlé nebo poškozené zboží dej do odpisu."),
        ("3", "Hotovost", "Řiď se pokynem aplikace. Pokud se vybírá, vlož vše do sáčku a napiš jen zobrazené ID."),
        ("4", "Kontrola", "Zkontroluj čistotu, funkci, ceny a zavření automatu. Pak dokonči práci na automatu."),
    ])
    add_callout(doc, "Expirace", "Starší zboží patří dopředu, novější dozadu. Prošlé zboží se nesmí vrátit mezi prodejné zásoby auta.", "warn")

    page_break(doc)
    add_heading(doc, "7. Návrat a ukončení směny", "Směna končí až po návratu, vyložení a zápisu konečných kilometrů.")
    add_two_column_visual(doc, images["today"], [
        ("1", "Dokonči poslední zastávku", "Na Trase musí být všechny navštívené zastávky uzavřené."),
        ("2", "Vylož vratky a odpisy", "Použitelné zboží vrať do skladu. Prošlé, rozbité nebo vysypané dej do odpisu podle aplikace."),
        ("3", "Zapiš konečný stav km", "Opiš skutečný tachometr po návratu. Nepiš počet ujetých km, ale celý konečný stav."),
        ("4", "Ukonči směnu", "Tlačítko použij až po dokončení všech povinných úkolů. Zkontroluj zelené potvrzení."),
    ])
    add_callout(doc, "Když aplikace nepustí dál", "Přečti červenou nebo oranžovou hlášku. Obvykle chybí potvrzené převzetí zboží, expirace, hotovost, kontrola automatu, inventura nebo konečné km.", "info")
    add_heading(doc, "Tři pravidla, která tě ochrání")
    add_step(doc, 1, "Potvrzuji jen skutečnost", "Co fyzicky nevidím nebo neudělám, to v aplikaci nepotvrdím.")
    add_step(doc, 2, "Nic neodhaduji", "Kusy, expirace a kilometry opisuji z fyzického stavu.")
    add_step(doc, 3, "Při problému zastavím", "Udělám snímek obrazovky a zavolám vedoucímu dřív, než pokračuji.")

    core = doc.core_properties
    core.title = "OLVEND - první směna Kristýny Dvořákové"
    core.subject = "Jednoduchý obrazový návod k mobilní aplikaci OLVEND"
    core.author = "OLMIKA"
    core.keywords = "OLVEND, mobilní aplikace, první směna, operátorka, FEFO"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
